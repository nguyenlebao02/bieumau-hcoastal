"""Fill Coastal template DOCX by replacing dotted runs with form data."""
import sys, json, re
from docx import Document

TEMPLATE = sys.argv[1]
OUTPUT = sys.argv[2]
data = json.loads(sys.argv[3])

FIELD_MAP = {
    3: ["csbhDate"],
    4: ["tenNhaKetNoi"],
    7: ["tenKhachHang"],
    8: ["giayToSo"],
    9: ["ngayCap", "noiCap"],
    10: ["diaChiHoKhau"],
    11: ["diaChiLienHe"],
    12: ["soDienThoai", "email"],
    14: ["maCan", "khuBietThu"],
    15: ["huong", "mauNha"],
    16: ["dienTichDat", "tongDienTichXayDung"],
    17: ["tongGiaXayDung", "chietKhau"],
    18: ["tongGiaBan"],
    22: ["chuKyNhay"],
    23: ["chuKyChinh"],
    25: ["ngayKy", "thangKy"],
}

DOT_PATTERN = re.compile(r'[….]+')  # Unicode ellipsis + ascii dot

doc = Document(TEMPLATE)

for pi, field_names in FIELD_MAP.items():
    if pi >= len(doc.paragraphs):
        continue
    para = doc.paragraphs[pi]
    
    # Find all dot position + length in combined run text
    segments = []  # (run_index, start_in_run, length)
    for ri, run in enumerate(para.runs):
        if not run.text:
            continue
        for m in DOT_PATTERN.finditer(run.text):
            segments.append((ri, m.start(), m.end() - m.start()))
    
    if not segments:
        continue
    
    # Group consecutive dot segments (within same or adjacent runs)
    groups = []
    current = []
    for si, (ri, start, length) in enumerate(segments):
        if current and ri == segments[current[-1]][0] + 1:
            # Adjacent run
            current.append(si)
        elif current and ri > segments[current[-1]][0] + 1:
            # Gap - new group
            groups.append(current)
            current = [si]
        else:
            current.append(si)
    if current:
        groups.append(current)
    
    # Fill fields left to right
    for fi, field_name in enumerate(field_names):
        value = str(data.get(field_name, "") or "")
        if not value:
            continue
        if fi >= len(groups):
            continue
        
        group = groups[fi]
        # Clear dots from all segments in this group
        for si in group:
            ri, start, length = segments[si]
            run = para.runs[ri]
            run.text = run.text[:start] + run.text[start+length:]
        
        # Insert value at the first segment's position
        first_si = group[0]
        ri, start, _ = segments[first_si]
        run = para.runs[ri]
        run.text = run.text[:start] + value + run.text[start:]

# Handle tables
if doc.tables:
    t = doc.tables[0]
    # Row 1: VVNH
    if len(t.rows) > 1:
        vvnh = data.get("vvnhLaiSuat", "Không")
        cell = t.rows[1].cells[1]
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""
        if cell.paragraphs:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = vvnh
            else:
                cell.paragraphs[0].add_run(vvnh)
    
    # Row 2: Quà tặng
    if len(t.rows) > 2:
        quaTang = data.get("quaTang", "")
        if quaTang:
            cell = t.rows[2].cells[0]
            for p in cell.paragraphs:
                for r in p.runs:
                    # Only replace pure-dot runs, not mixed content
                    if r.text and all(c in '…. ' for c in r.text) and any(c in '….' for c in r.text):
                        r.text = quaTang
                        break

doc.save(OUTPUT)
print(OUTPUT)
